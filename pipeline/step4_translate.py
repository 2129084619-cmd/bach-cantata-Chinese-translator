# -*- coding: utf-8 -*-
"""Step 4: Prepare translation context and generate Docx 2 (German-Chinese parallel).

This module:
1. Assembles all translation context into a structured JSON
2. Generates a Docx 2 with paragraph-format German-Chinese parallel text,
   endnotes with hyperlinks, and a glossary table.

The actual translation is performed by the AI assistant, which reads the
context and fills in the Chinese paragraphs.
"""

import json
import os
import re
import shutil
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import config
from .logger import get_logger

log = get_logger()


# ═══════════════════════════════════════════════════════════════
# ENDNOTE HYPERLINK INFRASTRUCTURE
# ═══════════════════════════════════════════════════════════════

# Counter for unique bookmark IDs
_bookmark_counter = [0]


def _next_bm_id():
    _bookmark_counter[0] += 1
    return str(_bookmark_counter[0])


def _add_bookmark(paragraph, name, bm_id=None):
    """Add a Word bookmark to a paragraph."""
    if bm_id is None:
        bm_id = _next_bm_id()
    p = paragraph._p
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bm_id)
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bm_id)
    p.insert(0, start)
    p.append(end)


def _add_hyperlink(paragraph, text, bookmark_name, superscript=True):
    """Add a clickable internal hyperlink to a paragraph.

    Args:
        paragraph: docx paragraph object
        text: display text (e.g., "[1]")
        bookmark_name: target bookmark name (e.g., "fn_1")
        superscript: render as superscript
    """
    p = paragraph._p
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), bookmark_name)
    hl.set(qn('w:history'), '1')

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    if superscript:
        vertAlign = OxmlElement('w:vertAlign')
        vertAlign.set(qn('w:val'), 'superscript')
        rPr.append(vertAlign)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '16')  # 8pt superscript
    rPr.append(sz)

    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    hl.append(r)
    p.append(hl)


def inject_footnote_hyperlinks(docx_path, bwv_number, footnote_map):
    """Post-process: inject superscript footnote hyperlinks into an existing docx.

    Use this when the Docx2 was generated without hyperlinks (e.g., BWV numbers
    whose English text lacks [N] markers in the JSON API).

    Args:
        docx_path: path to the existing docx file
        bwv_number: int or str
        footnote_map: dict {(movement, line_index): [footnote_ids]}
            e.g. {(1, 0): [1, 2, 3], (2, 0): [4, 5]}
    """
    import re as _re  # local import to avoid shadowing outer re
    doc = Document(docx_path)

    # Scan paragraphs to build (movement, line_index) → paragraph_index mapping
    current_mv = 0
    line_idx_in_mv = 0
    cn_para_indices = []  # [(movement, line_idx, paragraph_index)]

    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        m = _re.match(r'Movement\s+(\d+)', txt)
        if m and p.style.name.startswith('Heading'):
            current_mv = int(m.group(1))
            line_idx_in_mv = 0
            continue
        if any('\u4e00' <= c <= '\u9fff' or '\u3000' <= c <= '\u303f' for c in (p.text or '')) and not _re.match(r'^\[\d+\]', txt):
            # This is a Chinese translation paragraph (not an endnote)
            cn_para_indices.append((current_mv, line_idx_in_mv, i))
            line_idx_in_mv += 1

    # Inject hyperlinks
    injected = 0
    for (mv, li, p_idx) in cn_para_indices:
        fn_ids = footnote_map.get((mv, li), [])
        if fn_ids:
            p = doc.paragraphs[p_idx]
            for fn_id in fn_ids:
                _inject_hl_to_para(p, fn_id, bwv_number)
            injected += 1

    doc.save(docx_path)
    log.info(f"[Step 4] Injected {injected} footnote hyperlink groups into {os.path.basename(docx_path)}")
    return injected


def _inject_hl_to_para(para, fn_id, bwv):
    """Inject a single superscript hyperlink into a paragraph using lxml."""
    p_elem = para._p
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), f'fn_{bwv}_{fn_id}')
    hl.set(qn('w:history'), '1')

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '16')
    rPr.append(sz)

    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = f'[{fn_id}]'
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    hl.append(r)
    p_elem.append(hl)


def translate_footnotes_in_docx(docx_path, footnotes_cn, disclaimer='\u5185\u5bb9\u4ec5\u4f9b\u53c2\u8003'):
    """Replace English footnote texts with Chinese translations in an existing docx.

    Per policy (2026-08-16): footnotes originate from bachcantatatexts.org and are
    annotation-only. Each translated footnote is suffixed with a disclaimer
    ("内容仅供参考" by default) to flag its reference-only status.

    Args:
        docx_path: path to the docx file (modified in-place)
        footnotes_cn: dict {footnote_number: chinese_text}
        disclaimer: str, appended to each translated footnote ('' to disable)
    """
    import re as _re
    suffix = f'\u3000\uff08{disclaimer}\uff09' if disclaimer else ''
    doc = Document(docx_path)
    replaced = 0
    for p in doc.paragraphs:
        txt = p.text.strip()
        m = _re.match(r'^\[(\d+)\]\s', txt)
        if m:
            fn_num = int(m.group(1))
            if fn_num in footnotes_cn:
                # Check if already Chinese (skip)
                if any('\u4e00' <= c <= '\u9fff' for c in txt):
                    continue
                for run in p.runs:
                    run.text = ''
                # Set bold number marker
                if p.runs:
                    p.runs[0].text = f'[{fn_num}] '
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(9)
                # Add translated text (+ disclaimer)
                if len(p.runs) > 1:
                    p.runs[1].text = footnotes_cn[fn_num] + suffix
                    p.runs[1].font.size = Pt(9)
                elif p.runs:
                    p.runs[0].text = f'[{fn_num}] {footnotes_cn[fn_num]}{suffix}'
                replaced += 1
    doc.save(docx_path)
    log.info(f"[Step 4] Translated {replaced} footnotes in {os.path.basename(docx_path)}")
    return replaced


# ═══════════════════════════════════════════════════════════════
# TRANSLATION CONTEXT ASSEMBLY
# ═══════════════════════════════════════════════════════════════

def _g_text(line):
    """Extract plain text from a German line (str or role-label dict)."""
    if isinstance(line, dict):
        return line.get('text', ''), line.get('line_is_role_label', False)
    return str(line) if line else '', False


def prepare_translation_context(bwv_number, movements, footnotes, glossary,
                                bible_cn, luther_verify, metadata,
                                title=''):
    """Assemble comprehensive translation context.

    Args:
        bwv_number: int or str
        movements: list from step1
        footnotes: dict from step1
        glossary: list from step2.5
        bible_cn: dict from step3
        luther_verify: dict from step2.5
        metadata: dict from step2
        title: str, the work_name from the JSON API (overall cantata title)

    Returns:
        dict: Structured translation context
    """
    lines_with_context = []
    for mv in movements:
        mv_num = mv['number']
        de_lines = mv.get('german', [])
        line_fn_ids = mv.get('line_footnote_ids', [])
        non_role_idx = 0
        for i, de_line in enumerate(de_lines):
            de_text, is_role = _g_text(de_line)
            # line_footnote_ids is aligned to lyric lines only (role labels
            # excluded). Role-label lines carry no footnote; lyric lines advance
            # the non_role_idx counter (fixes dialogue-cantata footnote offset).
            if is_role:
                this_line_fn_ids = []
            else:
                this_line_fn_ids = line_fn_ids[non_role_idx] if non_role_idx < len(line_fn_ids) else []
                non_role_idx += 1

            # Find relevant Chinese Bible passages for these footnotes
            relevant_bible = {}
            for ref_key, ref_data in bible_cn.items():
                for fn_id in this_line_fn_ids:
                    if fn_id in ref_data.get('footnote_ids', []):
                        relevant_bible[ref_key] = ref_data.get('verses_text', '')

            # NOTE: the `english` field is deliberately NOT included as translation
            # reference (policy 2026-08-16) — bachcantatatexts.org English text is
            # annotation-only. Translation relies on German + Chinese CUV only.
            lines_with_context.append({
                'movement': mv_num,
                'line_index': i,
                'german': de_text,
                'is_role_label': is_role,
                'footnote_ids': this_line_fn_ids,
                'relevant_bible_cn': relevant_bible,
            })

    context = {
        'bwv': int(bwv_number),
        'title': title,
        'materials': {
            'glossary': glossary,
            'bible_cn_summary': {
                ref: data.get('verses_text', '')
                for ref, data in bible_cn.items()
            },
            'luther_verify_summary': (
                f'{luther_verify.get("total_references", 0)} references verified, '
                f'source: {luther_verify.get("source", "")}'
            ),
        },
        'lines': lines_with_context,
        'translation_guidelines': [
            '【最高优先级】宗教专有名词、神学术语、圣经引用严格对齐中文和合本 (CUV) 译文，不得自创译法',
            '【和合本优先】当注释的释义与和合本经文不一致时，以和合本为准',
            '德语原文 + 中文和合本经文是唯二的核心依据；不使用英文译文作为翻译参考',
            '须补充的神学或文化背景说明（超出和合本范围）应以脚注或方括号标注，明确与和合本正文区分',
            '每行德语对应一行中文，保持诗歌分行与段落结构',
            '在准确传达语义的前提下，兼顾中文表达的流畅与简洁',
            '路德宗神学概念（如圣餐的"in, mit und unter"、因信称义等）需准确传达',
        ],
    }

    return context


# ═══════════════════════════════════════════════════════════════
# DOCX 2 GENERATION (PARAGRAPH FORMAT + ENDNOTES + GLOSSARY)
# ═══════════════════════════════════════════════════════════════

def generate_docx2(bwv_number, movements, footnotes, glossary,
                   movement_info, metadata, folder_path, title=''):
    """Generate combined Docx: basic info + German-Chinese parallel translation.

    FORMAT (v2.2):
      - Title + basic info table (merged from former Docx1)
      - Each movement: heading with vocal/instrumental info
      - Per-line: German (TNR regular) → Chinese (宋体 regular), paired
      - Dialogue cantata: role labels as bold italic labels
      - Pipe-separated duet text split by role
      - Endnotes with hyperlink jump from superscript markers

    Args:
        title: str, the work_name from the JSON API (overall cantata title)
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')

    _bookmark_counter[0] = 0

    # Build movement info lookup from bach-cantatas.com
    mv_info_map = {}
    for mi in movement_info:
        mv_info_map[mi.get('number', 0)] = mi

    # ── Helpers ──
    def _set_font(run, font_name='Times New Roman', east_asian=None, bold=None, italic=None):
        """Ensure consistent font settings on a run."""
        run.font.name = font_name
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if east_asian:
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), east_asian)

    def _h(text, level=1):
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            _set_font(r, 'Times New Roman', east_asian='\u5b8b\u4f53')

    def _p(text, bold=False, italic=False, size=None, color=None,
           align=None, sa=Pt(6)):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        _set_font(r, 'Times New Roman', east_asian='\u5b8b\u4f53')
        if size: r.font.size = size
        if color: r.font.color.rgb = color
        if align is not None: p.alignment = align
        p.paragraph_format.space_after = sa
        return p

    def _de_line_para(text):
        """German lyric line — Times New Roman, REGULAR weight (not bold)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(text)
        r.font.size = Pt(11)
        _set_font(r, 'Times New Roman', bold=False)
        return p

    def _role_line_para(role_name):
        """Dialogue role line — bold TNR role name + 宋体 CN placeholder on same line."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        # Role name in bold TNR
        r1 = p.add_run(role_name)
        r1.font.size = Pt(11)
        _set_font(r1, 'Times New Roman', bold=True)
        # Spacer
        r2 = p.add_run('  ')
        r2.font.size = Pt(11)
        # Chinese placeholder
        r3 = p.add_run('\u3010\u5f85\u7ffb\u8bd1\u3011')
        r3.font.size = Pt(11)
        _set_font(r3, 'Times New Roman', east_asian='\u5b8b\u4f53')
        return p

    def _cn_line_para():
        """Chinese translation paragraph — 宋体 regular, black."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run('\u3010\u5f85\u7ffb\u8bd1\u3011')
        r.font.size = Pt(11)
        _set_font(r, 'Times New Roman', east_asian='\u5b8b\u4f53')
        return p

    def _hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single', qn('w:sz'): '6',
            qn('w:space'): '1', qn('w:color'): 'AAAAAA'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _gap():
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # TITLE
    # ═══════════════════════════════════════════════════════════════
    # Use the work_name from the JSON API (overall cantata title),
    # NOT movements[0]['german'][0] — the first movement may be an
    # instrumental sinfonia/concerto whose text line is just "Sinfonia".
    display_title = title
    if not display_title:
        # Fallback 1: chorale text from metadata (accurate, no trailing comma)
        display_title = (metadata or {}).get('chorale_text', '')
    if not display_title and movements and movements[0].get('german'):
        first = movements[0]['german'][0]
        # german lines may be dicts ({'text':..., 'is_chorale':True}) since v1.0.4
        display_title = first.get('text', '') if isinstance(first, dict) else str(first)
    _p(f'BWV {bwv_number} \u2014 \u201e{display_title}\u201c',
       bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(2))
    _p('\u5fb7\u4e2d\u5bf9\u7167\u8bd1\u6587 / German\u2013Chinese Parallel Translation',
       size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    _p('\u6570\u636e\u6765\u6e90\uff1aBachCantataTexts.org  &  Bach-Cantatas.com',
       italic=True, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
       align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(16))
    _hr()
    _gap()

    # ═══════════════════════════════════════════════════════════════
    # BASIC INFORMATION TABLE (merged from former Docx1)
    # ═══════════════════════════════════════════════════════════════
    _h('\u57fa\u672c\u4fe1\u606f / General Information', level=2)

    info_items = [
        ('\u4f5c\u54c1\u7f16\u53f7 / BWV', str(bwv_number)),
    ]
    occasion = metadata.get('occasion', '')
    if occasion:
        info_items.append(('\u793c\u4eea\u573a\u5408 / Occasion', occasion))
    composed = metadata.get('composed', '')
    if composed:
        info_items.append(('\u4f5c\u66f2\u65f6\u95f4 / Composed', composed))
    librettist = metadata.get('librettist', '')
    if librettist:
        info_items.append(('\u6b4c\u8bcd\u4f5c\u8005 / Librettist', librettist))
    chorale = metadata.get('chorale_text', '')
    if chorale:
        info_items.append(('\u4f17\u8d5e\u6b4c / Chorale', chorale))
    ep = metadata.get('readings', {}).get('epistle', {})
    if ep:
        info_items.append(('\u4e66\u4fe1\u7ecf\u6587 / Epistle',
                          f"{ep.get('book', '')} {ep.get('chapter', '')}:{ep.get('verses', '')}"))
    gos = metadata.get('readings', {}).get('gospel', {})
    if gos:
        info_items.append(('\u798f\u97f3\u7ecf\u6587 / Gospel',
                          f"{gos.get('book', '')} {gos.get('chapter', '')}:{gos.get('verses', '')}"))

    info_table = doc.add_table(rows=len(info_items), cols=2, style='Light Grid Accent 1')
    for i, (k, v) in enumerate(info_items):
        info_table.cell(i, 0).text = k
        info_table.cell(i, 1).text = v
        for ci in [0, 1]:
            for para in info_table.cell(i, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    for row in info_table.rows:
        row.cells[0].width = Cm(4.5)

    _gap()
    _hr()
    _gap()

    # ═══════════════════════════════════════════════════════════════
    # MOVEMENTS — per-line German-Chinese pairing
    # ═══════════════════════════════════════════════════════════════
    for mv in movements:
        num = int(mv.get('number', 0))
        de_lines = mv.get('german', [])
        en_lines = mv.get('english', [])
        line_fn_ids = mv.get('line_footnote_ids', [])
        is_role_list = mv.get('line_is_role_label', [])
        role_name_list = mv.get('line_role_name', [])
        is_duet_list = mv.get('line_is_duet', [])
        role_texts_list = mv.get('line_role_texts', [])

        if len(line_fn_ids) < len(de_lines):
            line_fn_ids = line_fn_ids + [[]] * (len(de_lines) - len(line_fn_ids))

        # ── Movement heading: include vocal and instrumental info ──
        mi = mv_info_map.get(num, {})
        raw_type = mi.get('type', '').strip()
        voices = mi.get('voices', '').strip()

        heading_parts = [f'Movement {num}']
        full_label = raw_type
        if voices and voices not in raw_type:
            full_label = f'{raw_type} [{voices}]'
        elif not full_label:
            full_label = mv.get('type', '')

        if full_label:
            heading_parts.append(full_label)

        _h(' \u2014 '.join(heading_parts), level=2)
        _gap()

        # ── Render each line ──
        non_role_idx = 0
        for idx, de_line in enumerate(de_lines):
            # Normalize: extract text from role-label dict (UAlberta step1 / step1.6)
            is_role = False
            if isinstance(de_line, dict):
                is_role = de_line.get('line_is_role_label', False)
                de_line = de_line.get('text', de_line.get('german', ''))
            # Fallback to is_role_list (bachcantatatexts.org parsing path)
            if not is_role:
                is_role = is_role_list[idx] if idx < len(is_role_list) else False
            role_name = role_name_list[idx] if idx < len(role_name_list) else None
            is_duet = is_duet_list[idx] if idx < len(is_duet_list) else False
            rt_map = role_texts_list[idx] if idx < len(role_texts_list) else None
            en_line = en_lines[idx] if idx < len(en_lines) else ''
            # line_footnote_ids is aligned to lyric lines only; role labels skip.
            if is_role:
                fn_ids = []
            else:
                fn_ids = line_fn_ids[non_role_idx] if non_role_idx < len(line_fn_ids) else []
                non_role_idx += 1

            if is_role:
                # Dialogue role — bold TNR role name + 宋体 CN on same line
                _role_line_para(de_line)

            elif is_duet and rt_map:
                # Pipe-separated duet line — split by role, render each part
                de_parts = rt_map.get('de', {})
                part_keys = sorted(de_parts.keys())
                for pk in part_keys:
                    dp = de_parts[pk]
                    # German text for this part with indent
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.left_indent = Cm(0.8)
                    r = p.add_run(dp)
                    r.font.size = Pt(11)
                    _set_font(r, 'Times New Roman', bold=False)
                    # Chinese placeholder
                    _cn_line_para()

            else:
                # Normal line (lyrics or chorale)
                # German line
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.5)
                r = p.add_run(de_line)
                r.font.size = Pt(11)
                _set_font(r, 'Times New Roman', bold=False)

                # Footnote superscript hyperlinks
                if fn_ids:
                    r_space = p.add_run(' ')
                    r_space.font.size = Pt(11)
                    for fn_id in fn_ids:
                        bookmark_name = f'fn_{bwv_number}_{fn_id}'
                        _add_hyperlink(p, f'[{fn_id}]', bookmark_name, superscript=True)

                # Chinese translation placeholder
                _cn_line_para()

        _gap()
        _gap()

    _hr()

    # ═══════════════════════════════════════════════════════════════
    # ENDNOTES
    # ═══════════════════════════════════════════════════════════════
    _h('\u5b66\u672f\u6ce8\u91ca / Scholarly Endnotes', level=1)

    if footnotes:
        for num_str in sorted(footnotes.keys(), key=int):
            n = int(num_str)
            text = footnotes[num_str]

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.3)

            bookmark_name = f'fn_{bwv_number}_{n}'
            _add_bookmark(p, bookmark_name)

            rn = p.add_run(f'[{n}] ')
            rn.bold = True
            rn.font.size = Pt(9)

            rt = p.add_run(text)
            rt.font.size = Pt(9)

            r_back = p.add_run(' \u2191')
            r_back.font.size = Pt(8)
            r_back.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    _gap()

    # ═══════════════════════════════════════════════════════════════
    # FOOTER
    # ═══════════════════════════════════════════════════════════════
    _hr()
    _p(f'\u672c\u6587\u6863\u7531 Bach Cantata Pipeline \u4e8e '
       f'{datetime.now().strftime("%Y-%m-%d")} \u81ea\u52a8\u751f\u6210\uff0c'
       f'\u4e2d\u6587\u7ffb\u8bd1\u5f85\u5b8c\u6210\u3002',
       italic=True, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
       align=WD_ALIGN_PARAGRAPH.CENTER)

    output_path = os.path.join(
        folder_path,
        f'BWV{bwv_number}_\u5fb7\u4e2d\u5bf9\u7167\u8bd1\u6587.docx'
    )
    # Keep any previously COMPLETED translation before overwriting (versioning).
    _archive_existing_docx(output_path)
    doc.save(output_path)
    log.info(f"[Step 4] Saved Docx: {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ═══════════════════════════════════════════════════════════════

def run(bwv_number, movements, footnotes, glossary, bible_cn,
        luther_verify, metadata, folder_path, title=''):
    """Execute Step 4: prepare context and generate docx2.

    Args:
        title: str, the work_name from the JSON API (overall cantata title)

    Returns:
        dict: {'context': ..., 'docx2_path': ...}
    """
    bwv = str(bwv_number)

    # 1. Prepare translation context
    context = prepare_translation_context(
        bwv, movements, footnotes, glossary, bible_cn, luther_verify, metadata,
        title=title
    )

    # 2. Save context as JSON
    data_dir = os.path.join(folder_path, 'data')
    context_path = os.path.join(data_dir, 'translation_context.json')
    with open(context_path, 'w', encoding='utf-8') as f:
        json.dump(context, f, ensure_ascii=False, indent=2)
    log.info(f"[Step 4] Saved translation_context.json ({len(context['lines'])} lines)")

    # 3. Generate Docx 2
    movement_info = metadata.get('movement_info', [])
    docx2_path = generate_docx2(
        bwv, movements, footnotes, glossary, movement_info, metadata, folder_path,
        title=title
    )

    return {
        'context': context,
        'docx2_path': docx2_path,
    }


def check_untranslated_footnotes(docx_path):
    """Check if a docx has untranslated English footnotes.

    Returns:
        (untranslated_count, total_footnote_paragraphs)
        Returns (0, 0) if no footnote section found.
    """
    doc = Document(docx_path)
    paras = [p.text.strip() for p in doc.paragraphs]

    # Find footnote section
    fn_start = None
    for i, pt in enumerate(paras):
        if '学术注释' in pt or 'Endnotes' in pt or re.search(r'(Notes|Annotation|Footnot)', pt):
            fn_start = i
            break

    if fn_start is None:
        return (0, 0)

    untranslated = 0
    total = 0
    for j in range(fn_start + 1, len(paras)):
        pj = paras[j]
        if not pj:
            continue
        # Check if this looks like a footnote paragraph
        # (starts with [N] or contains reference markers)
        if re.match(r'^\[\d+\]', pj) or (
            j > fn_start + 1 and not any('\u4e00' <= c <= '\u9fff' for c in pj)
            and len(pj) > 30
        ):
            total += 1
            if not any('\u4e00' <= c <= '\u9fff' for c in pj):
                untranslated += 1

    return (untranslated, total)


# ═══════════════════════════════════════════════════════════════
# Translation write-back with 1:1 alignment enforcement
# ═══════════════════════════════════════════════════════════════

def write_cantata_translations(docx_path, translations):
    """Write Chinese translations into a BWV cantata docx.

    Replaces 【待翻译】 placeholders sequentially with the provided
    translations. Enforces strict 1:1 alignment: if the number of
    translations doesn't match the number of placeholders, pads or
    truncates with a warning.

    Font conventions: Times New Roman 11pt for body, 宋体 for CJK.

    Args:
        docx_path: str — path to BWV{N}_德中对照译文.docx
        translations: list of str — one per placeholder, in docx order

    Returns:
        int: number of placeholders replaced
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(docx_path)
    total_ph = sum(1 for p in doc.paragraphs if '【待翻译】' in p.text)

    if len(translations) != total_ph:
        print(f"  [WARN] write_cantata_translations: {len(translations)} translations "
              f"but {total_ph} placeholders — padding/truncating to match")
        if len(translations) < total_ph:
            translations = list(translations) + ['[行缺]'] * (total_ph - len(translations))
        else:
            translations = translations[:total_ph]

    def _style_cn_run(r):
        r.font.size = Pt(11)
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.bold = False
        r.italic = False
        rPr = r._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')

    replaced = 0
    for p in doc.paragraphs:
        if '【待翻译】' in p.text and replaced < len(translations):
            cn_text = translations[replaced]

            # Detect a dialogue role line: it has a bold role-name run (e.g. "Seele",
            # "Bass") followed by the 【待翻译】 placeholder. Keep the bold role name
            # and replace only the placeholder run.
            has_bold_role = any(
                run.bold and run.text.strip() and '【待翻译】' not in run.text
                for run in p.runs
            )

            if has_bold_role:
                for run in p.runs:
                    if '【待翻译】' in run.text:
                        run.text = cn_text
                        _style_cn_run(run)
                        break
            else:
                for run in p.runs:
                    run.text = ''
                if p.runs:
                    _style_cn_run(p.runs[0])
                    p.runs[0].text = cn_text
            replaced += 1

    doc.save(docx_path)
    return replaced


# ═══════════════════════════════════════════════════════════════
# OUTPUT LOCATION HELPERS (2026-08-16)
#   - final translations live under "raw data & all translations/BWV_N/"
#   - "latest translations/BWV_N/" mirrors the newest docx + txt
#   - re-running never destroys a prior completed docx (timestamped archive)
# ═══════════════════════════════════════════════════════════════

def _archive_existing_docx(output_path):
    """Archive a previously completed docx before overwriting it.

    Preserves the prior translation as a timestamped copy in the same folder,
    so re-running the pipeline never silently destroys existing work.
    """
    if not os.path.exists(output_path):
        return
    # Timestamp = the file's own last-modified time (mtime), NOT the re-run date,
    # so the archive faithfully records when the prior translation was last touched.
    ts = datetime.fromtimestamp(os.path.getmtime(output_path)).strftime('%Y%m%d_%H%M%S')
    base, ext = os.path.splitext(output_path)
    backup_path = f'{base}_{ts}{ext}'
    shutil.copy2(output_path, backup_path)
    log.info(f"[Step 4] Archived prior docx → {os.path.basename(backup_path)}")


def mirror_to_latest(bwv_number, docx_path=None, txt_path=None):
    """Mirror the newest docx/txt into "latest translations/BWV_N/".

    This directory always reflects the most recent completed translation.
    """
    latest_dir = os.path.join(config.LATEST_TRANSLATIONS_DIR, f'BWV_{bwv_number}')
    os.makedirs(latest_dir, exist_ok=True)
    copied = []
    for src in (docx_path, txt_path):
        if src and os.path.exists(src):
            dst = os.path.join(latest_dir, os.path.basename(src))
            shutil.copy2(src, dst)
            copied.append(dst)
    if copied:
        log.info(f"[Step 4] Mirrored {len(copied)} file(s) → {latest_dir}")
    return copied


def write_chinese_txt(bwv_number, folder_path, lines, mirror=True):
    """Write the plain Chinese txt into the BWV folder (overwrites) and
    optionally mirror it to "latest translations".

    Args:
        bwv_number: int or str
        folder_path: str — the BWV folder under "raw data & all translations"
        lines: list of str — txt content lines
        mirror: bool — also copy into "latest translations/BWV_N/"
    """
    txt_path = os.path.join(folder_path, f'BWV{bwv_number}_中文译文.txt')
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    if mirror:
        mirror_to_latest(bwv_number, txt_path=txt_path)
    return txt_path
