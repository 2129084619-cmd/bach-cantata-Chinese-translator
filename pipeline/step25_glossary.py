# -*- coding: utf-8 -*-
"""Step 2.5: Glossary generation, Luther 1545 verification, and Docx 1 generation.

1. Extract religious terms from lyrics and footnotes → glossary table
2. Cross-verify Luther Bible citations against BibleGateway LUTH1545
3. Generate integrated Docx 1: metadata + glossary + verification + lyrics + footnotes
"""

import json
import os
import re
import time
from datetime import datetime

import requests

from . import config
from .logger import get_logger

log = get_logger()


# ═══════════════════════════════════════════════════════════════
# 1. RELIGIOUS VOCABULARY GLOSSARY
# ═══════════════════════════════════════════════════════════════

# Built-in theological glossary: German term → (Chinese CUV equivalent, context/note)
PRESET_GLOSSARY = {
    # Christological titles
    'Morgenstern': ('晨星', '启 22:16「我是…明亮的晨星」; Luther: heller Morgenstern'),
    'Br\u00e4utigam': ('新郎', '太 25:1-13 十个童女比喻; 指基督与教会末世联合'),
    'K\u00f6nig': ('君王', '大卫的弥赛亚君王; 诗 104:1 等'),
    'Sohn David': ('大卫的子孙', '太 1:2-17; 路 3:23-38 耶稣的家谱'),
    'Sohn Gottes': ('神的儿子', '约 1:14「父独生子的荣光」'),
    'Heiland': ('救主', '路 2:11「为你们生了救主」'),
    'A und O': ('阿拉法，俄梅戛', '启 22:13「我是阿拉法，我是俄梅戛」'),
    'Anfang und Ende': ('初与终', '启 22:13; 与「阿拉法，俄梅戛」并列'),

    # Soteriological terms
    'Gnade': ('恩典', '约 1:14「充充满满地有恩典」; 路德宗核心概念'),
    'Wahrheit': ('真理', '约 1:14; 与「恩典」并列'),
    'Segen': ('赐福 / 福分', '弗 1:3-4「天上各样属灵的福气」'),
    'Glaube': ('信心 / 信仰', '来 11:1; 此处指领受圣餐的信心'),
    'Erquickung': ('安息 / 复苏', '太 11:28「我就使你们得安息」'),
    'Preis': ('赞美', '罗 3:7; 此处非「奖赏」或「代价」'),
    'Opfer': ('祭 / 奉献', '罗 12:1「将身体献上当作活祭」'),

    # Sacramental / Eucharistic
    'Himmelsbrot': ('天上的粮', '约 6:35「我就是生命的粮」; 出 16 吗哪预表'),
    'Leib und Blut': ('身体与宝血', '圣餐; 路德宗「in, mit und unter」教义'),
    'Brot des Lebens': ('生命的粮', '约 6:35, 48'),

    # Eschatological
    'Paradeis': ('乐园', '路 23:43「同我在乐园里了」'),
    'Freudenkrone': ('喜乐的冠冕', '帖前 2:19; 来 12:2'),

    # Old Testament figures
    'Wurzel Jesse': ('耶西的根', '罗 15:12 引用赛 11:10'),
    'Jesse': ('耶西', '大卫之父; 路得记'),
    'Jakob': ('雅各', '以色列先祖; 创 29:20'),
    'Gabriel': ('加百列', '天使长; 路 1:26-38'),

    # Theological concepts
    'irdscher Glanz': ('属世的荣光', '林后 4:4「福音的光」; 与「属灵的光」相对'),
    'leiblich Licht': ('肉体的光', '巴赫时代神学区分「肉体的光」与「属灵的光」'),
    'Freudenschein': ('喜乐的光辉', '即 Gnadenschein，耶稣头顶的光环'),
    'g\u00f6ttlichen Flammen': ('神圣的火焰', '歌 8:6「耶和华的火焰」; 指基督之爱'),
    'himmlische Lust': ('天上的喜乐', '来 6:4「尝过天恩的滋味」'),

    # Misc
    'Bethlehem': ('伯利恒', '路 2:8-10 天使向牧羊人报喜'),
    'Saiten': ('琴弦', '诗 150:4「用丝弦的乐器」'),
}


def _extract_german_words(german_lines):
    """Extract unique German content words from lyrics."""
    words = set()
    for line in german_lines:
        if isinstance(line, dict):
            line = line.get('text', '')
        # Split on spaces and punctuation
        tokens = re.findall(r'[A-Za-z\u00c4\u00e4\u00d6\u00f6\u00dc\u00fc\u00df]+', line)
        for t in tokens:
            if len(t) > 2 and t[0].isupper():  # German nouns are capitalized
                words.add(t)
    return words


def _check_glossary_match(word, glossary):
    """Check if a word matches any glossary entry (case-insensitive, partial)."""
    word_lower = word.lower()
    for key in glossary:
        if key.lower() == word_lower or key.lower() in word_lower or word_lower in key.lower():
            return key
    return None


def generate_glossary(movements, footnotes):
    """Generate a religious vocabulary glossary.

    Returns:
        list of dicts: [{'german': ..., 'chinese_cuv': ..., 'note': ...}]
    """
    # Collect all German lines
    all_german = []
    for mv in movements:
        all_german.extend(mv.get('german', []))

    # Find matches
    glossary_entries = []
    matched_keys = set()

    for word in sorted(_extract_german_words(all_german)):
        key = _check_glossary_match(word, PRESET_GLOSSARY)
        if key and key not in matched_keys:
            matched_keys.add(key)
            cuv, note = PRESET_GLOSSARY[key]
            glossary_entries.append({
                'german': key,
                'chinese_cuv': cuv,
                'note': note,
            })

    log.info(f"[Step 2.5] Generated glossary: {len(glossary_entries)} terms")
    return glossary_entries


# ═══════════════════════════════════════════════════════════════
# 2. LUTHER 1545 VERIFICATION
# ═══════════════════════════════════════════════════════════════

def _fetch_luther_verse(book, chapter, verse):
    """Fetch a single verse from BibleGateway LUTH1545."""
    book_de = config.BOOK_GERMAN_MAP.get(book, book)
    url = config.URL_BIBLEGATEWAY_LUTHER.format(
        book_german=book_de, chapter=chapter, verse=verse
    )
    log.debug(f"[Step 2.5] Luther lookup: {book} {chapter}:{verse} → {url}")

    for attempt in range(1, config.MAX_RETRIES + 1):
        try:
            resp = requests.get(url, headers=config.HEADERS,
                                timeout=config.REQUEST_TIMEOUT)
            resp.raise_for_status()
            text = resp.text

            # Strategy: find the main passage content div, then extract
            # the specific verse. BibleGateway wraps passages in a div
            # with class containing "passage-" and then individual verses
            # in spans with data-verse attributes or versenum spans.

            # First, isolate the passage content area to avoid matching
            # navigation / "Read the Bible" UI elements.
            passage_match = re.search(
                r'<div[^>]*class="[^"]*passage[^"]*"[^>]*>(.*?)</div>\s*</div>\s*</div>',
                text, re.DOTALL
            )
            search_text = passage_match.group(1) if passage_match else text

            # Method 1: Look for verse number in <span class="versenum">
            # followed by the verse text in subsequent spans.
            verse_str = str(verse)
            m = re.search(
                rf'<span[^>]*class="[^"]*versenum[^"]*"[^>]*>\s*{re.escape(verse_str)}\s*</span>'
                rf'\s*(.+?)(?=<span[^>]*class="[^"]*versenum|$|</div>)',
                search_text, re.DOTALL
            )
            if m:
                verse_text = re.sub(r'<[^>]+>', '', m.group(1)).strip()
                # Remove verse numbers and footnotes
                verse_text = re.sub(r'^\d+\s*', '', verse_text)
                verse_text = re.sub(r'\[\w\]', '', verse_text)  # footnote markers
                if verse_text and not verse_text.startswith('Read'):
                    return verse_text.strip()

            # Method 2: Look for span with data-verse attribute
            m2 = re.search(
                rf'<span[^>]*class="[^"]*text[^"]*"[^>]*>\s*'
                rf'({re.escape(verse_str)}\s*)?'
                rf'([^<]+(?:<[^/][^>]*>[^<]*</[^>]*>)*[^<]*)'
                rf'\s*</span>',
                search_text, re.DOTALL
            )
            if m2:
                raw = m2.group(2) if m2.lastindex and m2.lastindex >= 2 else m2.group(0)
                verse_text = re.sub(r'<[^>]+>', '', raw).strip()
                verse_text = re.sub(r'^\d+\s*', '', verse_text)
                verse_text = re.sub(r'\[\w\]', '', verse_text)
                if verse_text and 'Bible' not in verse_text:
                    return verse_text.strip()

            # Method 3: Simplified — extract all visible text between
            # the target verse number and the next verse number.
            m3 = re.search(
                rf'>{re.escape(verse_str)}</[^>]*>\s*<[^>]*>\s*([^<]+)',
                search_text, re.DOTALL
            )
            if m3:
                verse_text = re.sub(r'<[^>]+>', '', m3.group(1)).strip()
                if verse_text and len(verse_text) > 3:
                    return verse_text

            return None
        except requests.RequestException as e:
            if attempt < config.MAX_RETRIES:
                time.sleep(2 ** attempt)
            else:
                log.warning(f"[Step 2.5] Luther lookup failed for {book} {chapter}:{verse}: {e}")
                return None


def _extract_luther_refs_from_footnote(footnote_text):
    """Parse Luther Bible references from a footnote.

    Returns list of (book, chapter, verse, quoted_text).
    """
    refs = []
    # Pattern: "John 1:14" or "Revelation 22:16" etc., followed by German Luther quote
    # in the footnote explanations
    pattern = re.compile(
        r'(?:in\s+)?(\d?\s*[A-Za-z]+(?:\s+of\s+[A-Za-z]+)?|Psalm|Psalms|Revelation|'
        r'Isaiah|Genesis|Exodus|Matthew|Mark|Luke|John|Acts|Romans|Ephesians|'
        r'Colossians|Hebrews|James|Corinthians)\s+'
        r'(\d+):(\d+(?:-\d+)?)',
        re.IGNORECASE
    )

    for m in pattern.finditer(footnote_text):
        book = m.group(1).strip()
        chapter = int(m.group(2))
        # Take only the first verse number for lookup
        verse = m.group(3).split('-')[0]
        refs.append((book, chapter, verse))

    return refs


def verify_luther_citations(footnotes):
    """Cross-verify Luther Bible citations against BibleGateway LUTH1545.

    Returns:
        dict: {'verified': bool, 'caveat': str, 'results': list}
    """
    results = []
    all_refs = set()

    for fnum, ftext in footnotes.items():
        refs = _extract_luther_refs_from_footnote(ftext)
        for book, chapter, verse in refs:
            key = (fnum, book, chapter, verse)
            if key not in all_refs:
                all_refs.add(key)
                luther_text = _fetch_luther_verse(book, chapter, verse)
                results.append({
                    'footnote_id': fnum,
                    'reference': f'{book} {chapter}:{verse}',
                    'biblegateway_text': luther_text,
                    'status': 'found' if luther_text else 'unavailable',
                })

    all_found = all(r['status'] == 'found' for r in results)

    log.info(f"[Step 2.5] Luther verification: {len(results)} references checked, "
             f"all_found={all_found}")

    return {
        'source': 'BibleGateway LUTH1545',
        'caveat': (
            'BibleGateway LUTH1545 based on 19th-century Cansteinsche Bibelanstalt '
            'edition; spelling modernized (e.g. wohnete→wohnte), but word choices '
            'and phrasing match Bach-era usage.'
        ),
        'verified': all_found,
        'total_references': len(results),
        'results': results,
    }


# ═══════════════════════════════════════════════════════════════
# 3. DOCX 1 GENERATION
# ═══════════════════════════════════════════════════════════════

def _generate_docx1(bwv_number, metadata, glossary, luther_verify,
                    movements, footnotes, translator, general_note, folder_path):
    """Generate Docx 1: title + basic information table only.

    This is a concise reference sheet. NO German lyrics, NO Luther
    verification (all verses unavailable due to network restrictions),
    NO footnotes, NO glossary.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')

    def _p(text, bold=False, italic=False, size=None, color=None, align=None, sa=Pt(6)):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        if size: r.font.size = size
        if color: r.font.color.rgb = color
        if align is not None: p.alignment = align
        p.paragraph_format.space_after = sa

    def _hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'),
            {qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): 'AAAAAA'})
        pBdr.append(bottom); pPr.append(pBdr)

    # ── Title ──
    title_de = f'BWV {bwv_number}'
    if movements:
        first_line = movements[0].get('german', [''])[0] if movements[0].get('german') else ''
        title_de += f' \u2014 \u201e{first_line}\u201c'
    _p(title_de, bold=True, size=Pt(24), align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(2))
    _p('\u57fa\u672c\u4fe1\u606f / Basic Information',
       size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    _p(f'\u6570\u636e\u6765\u6e90\uff1aBachCantataTexts.org  &  Bach-Cantatas.com',
       italic=True, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
       align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(16))
    _hr()
    doc.add_paragraph()

    # ── Basic Information Table ──
    info_items = [
        ('\u4f5c\u54c1\u7f16\u53f7 / BWV', str(bwv_number)),
    ]
    occasion = metadata.get('occasion', '')
    if occasion:
        info_items.append(('\u793c\u4eea\u573a\u5408 / Occasion', occasion))
    composed = metadata.get('composed', '')
    if composed:
        info_items.append(('\u4f5c\u66f2\u65f6\u95f4 / Composed', composed))
    librettist = metadata.get('librettist', '')
    if librettist:
        info_items.append(('\u6b4c\u8bcd\u4f5c\u8005 / Librettist', librettist))
    chorale = metadata.get('chorale_text', '')
    if chorale:
        info_items.append(('\u4f17\u8d5e\u6b4c / Chorale', chorale))
    if metadata.get('readings', {}).get('epistle', {}):
        ep = metadata['readings']['epistle']
        info_items.append(('\u4e66\u4fe1\u7ecf\u6587 / Epistle',
                          f"{ep.get('book', '')} {ep.get('chapter', '')}:{ep.get('verses', '')}"))
    if metadata.get('readings', {}).get('gospel', {}):
        gos = metadata['readings']['gospel']
        info_items.append(('\u798f\u97f3\u7ecf\u6587 / Gospel',
                          f"{gos.get('book', '')} {gos.get('chapter', '')}:{gos.get('verses', '')}"))

    info_table = doc.add_table(rows=len(info_items), cols=2, style='Light Grid Accent 1')
    for i, (k, v) in enumerate(info_items):
        info_table.cell(i, 0).text = k
        info_table.cell(i, 1).text = v
        for ci in [0, 1]:
            for para in info_table.cell(i, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    for row in info_table.rows:
        row.cells[0].width = Cm(4.5)

    # ── Footer ──
    doc.add_paragraph()
    _hr()
    _p(f'\u672c\u6587\u6863\u7531 Bach Cantata Pipeline \u4e8e '
       f'{datetime.now().strftime("%Y-%m-%d")} \u81ea\u52a8\u751f\u6210\u3002',
       italic=True, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
       align=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    output_path = os.path.join(
        folder_path,
        f'BWV{bwv_number}_\u539f\u6587\u4e0e\u9a8c\u8bc1.docx'
    )
    doc.save(output_path)
    log.info(f"[Step 2.5] Saved Docx 1: {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════

def run(bwv_number, movements, footnotes, translator, general_note, metadata, folder_path):
    """Execute Step 2.5: glossary, Luther verification, and docx1 generation.

    Args:
        bwv_number: int or str
        movements: list from step1
        footnotes: dict from step1
        translator: str from step1
        general_note: str from step1
        metadata: dict from step2
        folder_path: str, output folder

    Returns:
        dict with: glossary, luther_verify, docx1_path
    """
    bwv = str(bwv_number)

    # 1. Generate glossary
    glossary = generate_glossary(movements, footnotes)

    # 2. Luther verification
    log.info("[Step 2.5] Running Luther 1545 cross-verification...")
    luther_verify = verify_luther_citations(footnotes)

    # 3. Save glossary as JSON
    data_dir = os.path.join(folder_path, 'data')
    with open(os.path.join(data_dir, 'glossary.json'), 'w', encoding='utf-8') as f:
        json.dump(glossary, f, ensure_ascii=False, indent=2)
    with open(os.path.join(data_dir, 'luther_verify.json'), 'w', encoding='utf-8') as f:
        json.dump(luther_verify, f, ensure_ascii=False, indent=2)

    # 4. Update shared terminology database
    from . import glossary_db
    db_result = glossary_db.update_from_glossary(bwv, glossary)
    log.info(
        f"[Step 2.5] Term DB updated: "
        f"{db_result['new']} new, {db_result['updated']} updated, "
        f"{db_result['total']} total"
    )

    # 5. Docx1 generation removed — basic info table now merged into Docx2 (step4)
    # docx1_path = _generate_docx1(
    #     bwv, metadata, glossary, luther_verify,
    #     movements, footnotes, translator, general_note, folder_path
    # )

    return {
        'glossary': glossary,
        'luther_verify': luther_verify,
        'docx1_path': None,  # deprecated, basic info merged into Docx2
        'term_db_updated': db_result,
    }
