# -*- coding: utf-8 -*-
"""Chorale API — clean programmatic interfaces for chorale operations.

Two primary interfaces:
  1. query_chorale_docx(bwv)    — input BWV, locate and open .docx file
  2. translate_chorale(bwv)     — load chorale data, present translation context,
                                  then write Chinese translations back to .docx
"""

import json
import os
import sys
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

from . import chorale_config as cfg
from . import chorale_index as idx
from . import chorale_scraper as scraper


# ═══════════════════════════════════════════════════════════════
# Interface 1: Query & Open Docx
# ═══════════════════════════════════════════════════════════════

def query_chorale_docx(bwv_number):
    """Input a BWV number, locate and open the corresponding chorale .docx file.

    Returns:
        dict: {
            'bwv': str,
            'chorales_found': int,
            'docs': [{'chorale_id': str, 'title': str, 'docx_path': str, 'exists': bool}]
        }
    """
    bwv_str = str(bwv_number)
    entries = idx.lookup_by_bwv(bwv_str)

    result = {
        'bwv': bwv_str,
        'chorales_found': len(entries),
        'docs': []
    }

    if not entries:
        return result

    for entry in entries:
        cid = entry.get('chorale_id') or entry.get('id')
        if not cid:
            continue

        docx_path = os.path.join(cfg.DOCX_DIR, f'{cid}_德中对照译文.docx')
        exists = os.path.exists(docx_path)

        doc_info = {
            'chorale_id': cid,
            'title': entry.get('title', cid),
            'docx_path': docx_path,
            'exists': exists,
        }
        result['docs'].append(doc_info)

        if exists:
            _open_file(docx_path)

    return result


def _open_file(path):
    """Open a file with the system default application."""
    if sys.platform == 'win32':
        os.startfile(os.path.normpath(path))
    elif sys.platform == 'darwin':
        import subprocess
        subprocess.run(['open', path])
    else:
        import subprocess
        subprocess.run(['xdg-open', path])


# ═══════════════════════════════════════════════════════════════
# Interface 2: Manual Translation Execution
# ═══════════════════════════════════════════════════════════════

def translate_chorale(bwv_number, overwrite=False):
    """Execute translation for chorales associated with a single BWV number.

    Workflow:
      1. Look up chorale(s) for the BWV in the index
      2. Load scraped JSON data (or scrape if not available)
      3. Ensure .docx with placeholders exists (generate if missing)
      4. Present full translation context to the AI assistant
      5. AI produces Chinese translations per verse
      6. Write translations back to .docx via write_translations()

    Args:
        bwv_number: int or str
        overwrite: if True, re-scrape detail page even if data exists

    Returns:
        dict with translation context, or None if no chorales found
    """
    bwv_str = str(bwv_number)
    entries = idx.lookup_by_bwv(bwv_str)

    if not entries:
        return None

    all_contexts = []

    for entry in entries:
        cid = entry.get('chorale_id') or entry.get('id')
        if not cid:
            continue

        # Step 1: Load or scrape data
        chorale_data = scraper.load_chorale_data(cid)
        if not chorale_data or overwrite:
            chorale_data = scraper.scrape_chorale_detail(cid)
            scraper.save_chorale_data(cid, chorale_data)

        # Step 2: Ensure docx exists
        docx_path = os.path.join(cfg.DOCX_DIR, f'{cid}_德中对照译文.docx')
        if not os.path.exists(docx_path) or overwrite:
            from . import chorale_translator
            chorale_translator.generate_chorale_docx(chorale_data, chorale_id=cid)

        # Step 3: Read current docx state — extract placeholders
        placeholder_count = _count_placeholders(docx_path)

        # Step 4: Build translation context
        context = _build_translation_context(chorale_data, cid, docx_path, placeholder_count)
        all_contexts.append(context)

    # Step 5: Present unified context to AI
    _present_translation_context(all_contexts, bwv_str)

    return {
        'bwv': bwv_str,
        'chorales': all_contexts,
        'total_placeholders': sum(c['placeholder_count'] for c in all_contexts),
    }


# ═══════════════════════════════════════════════════════════════
# Interface 2b: Full Translation Pipeline (detect → overwrite/new → present)
# ═══════════════════════════════════════════════════════════════

def run_translate_pipeline(bwv_number):
    """Execute the full chorale translation pipeline with detection and overwrite.

    Steps:
      1. [检测] Look up BWV → chorale(s) in index; report what exists
      2. [判断] Determine status per chorale:
           - FULL: docx has 0 placeholders AND JSON has chinese_text → overwrite
           - PARTIAL: docx has placeholders but JSON has some chinese_text → re-present
           - TEMPLATE: docx exists with placeholders, no chinese_text → continue
           - NEW: nothing exists → scrape + generate
      3. [覆盖/新建] Regenerate docx with fresh placeholders (overwrite old)
                      or scrape + generate (new)
      4. [准备] Present full translation context
      5. [返回] Return context dict for AI to complete translation + write-back

    Args:
        bwv_number: int or str — BWV number

    Returns:
        dict: {
            'bwv': str,
            'status': dict — per-chorale detection results,
            'contexts': list — translation contexts ready for AI,
            'total_placeholders': int,
        }
        Returns None if no chorales found for this BWV.
    """
    bwv_str = str(bwv_number)

    # ─── Step 1: Detect ───────────────────────────────────────
    print(f"\n{'─' * 60}")
    print(f"  [检测] 正在查找 BWV {bwv_str} 对应的众赞歌...")
    print(f"{'─' * 60}")

    entries = idx.lookup_by_bwv(bwv_str)
    if not entries:
        print(f"  [错误] 未在索引中找到 BWV {bwv_str} 对应的众赞歌")
        print(f"        请运行 python -m 巴赫康塔塔中的众赞歌.chorale_main --rebuild-index 续建索引")
        return None

    n_entries = len(entries)
    names = ', '.join(
        e.get('chorale_id', e.get('id', '?')) for e in entries
    )
    print(f"  ✓ 找到 {n_entries} 首众赞歌: {names}")
    print()

    # ─── Step 2: Assess ───────────────────────────────────────
    status_map = {}   # chorale_id → status dict
    pre_existing = []  # chorales that had translations before

    for entry in entries:
        cid = entry.get('chorale_id') or entry.get('id')
        if not cid:
            continue

        docx_path = os.path.join(cfg.DOCX_DIR, f'{cid}_德中对照译文.docx')
        json_path = os.path.join(cfg.DATA_DIR, f'{cid}.json')

        docx_exists = os.path.exists(docx_path)
        json_exists = os.path.exists(json_path)

        # Determine status
        json_has_cn = False
        n_cn_verses = 0
        n_total_verses = 0
        docx_placeholders = -1
        docx_is_translated = False

        if json_exists:
            with open(json_path, 'r', encoding='utf-8') as f:
                chorale_data = json.load(f)
            cn = chorale_data.get('chinese_text', {})
            german = chorale_data.get('german_text', {})
            n_cn_verses = len(cn)
            n_total_verses = len(german)
            json_has_cn = n_cn_verses > 0

        if docx_exists:
            docx_placeholders = _count_placeholders(docx_path)
            docx_is_translated = docx_placeholders == 0

        # Classify
        if docx_is_translated or json_has_cn:
            status = 'FULL'
            pre_existing.append(cid)
            detail = f'已翻译 {n_cn_verses}/{n_total_verses} 诗节'
        elif docx_exists and docx_placeholders > 0:
            status = 'TEMPLATE'
            detail = f'docx 模板已存在 ({docx_placeholders} 个占位符)'
        else:
            status = 'NEW'
            detail = '首次翻译'

        status_map[cid] = {
            'status': status,
            'detail': detail,
            'docx_path': docx_path,
            'json_path': json_path,
            'docx_exists': docx_exists,
            'json_has_cn': json_has_cn,
            'n_cn_verses': n_cn_verses,
            'n_total_verses': n_total_verses,
            'docx_placeholders': docx_placeholders,
            'docx_is_translated': docx_is_translated,
        }

        # Print detection result
        icon = {'FULL': '✓', 'TEMPLATE': '○', 'NEW': '＋'}[status]
        action = {
            'FULL': '[覆盖旧翻译]',
            'TEMPLATE': '[继续翻译]',
            'NEW': '[新建翻译]',
        }[status]
        print(f"  {icon} [{status:8s}] {cid:12s} {action} {detail}")

    # ─── Step 3: Overwrite / Generate ─────────────────────────
    print(f"\n{'─' * 60}")
    if pre_existing:
        print(f"  [覆盖] 重新生成 {len(pre_existing)} 首已有翻译的众赞歌 docx...")
    print(f"{'─' * 60}")

    all_contexts = []
    for entry in entries:
        cid = entry.get('chorale_id') or entry.get('id')
        if not cid:
            continue

        st = status_map[cid]
        mode = '覆盖旧翻译，重新生成' if st['status'] == 'FULL' else '新建，抓取数据'
        print(f"\n  → {cid}: {mode}...")

        # Step 3a: Always regenerate docx with fresh placeholders
        #           (scrape if needed, then generate)
        chorale_data = scraper.load_chorale_data(cid)
        if not chorale_data:
            print(f"    [抓取] 正在从网站抓取 {cid} 详情...")
            try:
                chorale_data = scraper.scrape_chorale_detail(cid)
                scraper.save_chorale_data(cid, chorale_data)
                print(f"    ✓ 抓取完成")
            except Exception as e:
                print(f"    [错误] 抓取 {cid} 失败: {e}")
                continue
        else:
            # Reload JSON to ensure we have latest
            json_path = os.path.join(cfg.DATA_DIR, f'{cid}.json')
            if os.path.exists(json_path):
                with open(json_path, 'r', encoding='utf-8') as f:
                    chorale_data = json.load(f)

        # Step 3b: Generate fresh docx (overwrites existing)
        from . import chorale_translator
        docx_path = chorale_translator.generate_chorale_docx(
            chorale_data, chorale_id=cid
        )
        placeholder_count = _count_placeholders(docx_path)
        total_ger_lines = sum(
            len(v.get('lines', []))
            for v in chorale_data.get('german_text', {}).values()
        )
        print(f"    ✓ docx 已生成: {placeholder_count}/{total_ger_lines} 占位符 (逐行德中对照)")

        # Step 3c: Build context
        ctx = _build_translation_context(
            chorale_data, cid, docx_path, placeholder_count
        )
        all_contexts.append(ctx)

    # ─── Step 4: Present ──────────────────────────────────────
    _present_translation_context(all_contexts, bwv_str)
    print(f"  [准备完成] 共 {sum(c['placeholder_count'] for c in all_contexts)} 个占位符，等待 AI 翻译并写回")

    return {
        'bwv': bwv_str,
        'status': status_map,
        'contexts': all_contexts,
        'total_placeholders': sum(c['placeholder_count'] for c in all_contexts),
    }


def _count_placeholders(docx_path):
    """Count 【待翻译】 placeholders in a docx file."""
    doc = Document(docx_path)
    count = 0
    for p in doc.paragraphs:
        if '【待翻译】' in p.text:
            count += 1
    return count


def _build_translation_context(chorale_data, chorale_id, docx_path, placeholder_count):
    """Build structured translation context for a single chorale."""
    german_text = chorale_data.get('german_text', {})
    english_text = chorale_data.get('english_text', {})
    bach_verses = chorale_data.get('bach_verses', [])

    # Flatten verses for translation: [(verse_num, is_bach, [german_lines], [english_lines])]
    verses = []
    for vnum in sorted(german_text.keys()):
        gv = german_text[vnum]
        ev = english_text.get(vnum, {})
        verses.append({
            'verse_num': vnum,
            'is_bach': gv.get('bold', False) or vnum in bach_verses,
            'german_lines': gv.get('lines', []),
            'english_lines': ev.get('lines', []),
        })

    return {
        'chorale_id': chorale_id,
        'title': chorale_data.get('title', ''),
        'author': chorale_data.get('author', ''),
        'author_year': chorale_data.get('author_year', ''),
        'melody': chorale_data.get('melody', ''),
        'composer': chorale_data.get('composer', ''),
        'composer_year': chorale_data.get('composer_year', ''),
        'ekg': chorale_data.get('ekg', ''),
        'translator': chorale_data.get('translator', ''),
        'docx_path': docx_path,
        'placeholder_count': placeholder_count,
        'verses': verses,
        'bach_verses': bach_verses,
        'source_url': chorale_data.get('source_url', ''),
    }


def _present_translation_context(all_contexts, bwv_str):
    """Present translation context to the AI assistant."""
    print(f"\n{'=' * 70}")
    print(f" 众赞歌翻译上下文 — BWV {bwv_str}")
    print(f"{'=' * 70}")
    print(f" 共 {len(all_contexts)} 首众赞歌待翻译")
    print()

    for ctx in all_contexts:
        title = ctx['title']
        cid = ctx['chorale_id']
        n_verses = len(ctx['verses'])
        n_ph = ctx['placeholder_count']
        bach = ctx['bach_verses']

        print(f" {'─' * 60}")
        print(f" {title}  ({cid})")
        print(f" 作者: {ctx['author']} ({ctx['author_year']})")
        print(f" 旋律: {ctx['melody']}")
        print(f" 巴赫采用诗节: {bach}")
        print(f" 总诗节: {n_verses} | 占位符: {n_ph}")
        print(f" Docx: {ctx['docx_path']}")
        print()

        for v in ctx['verses']:
            vnum = v['verse_num']
            marker = ' [BACH]' if v['is_bach'] else ''
            print(f"  {vnum}.{marker}")
            for line in v['german_lines']:
                print(f"    {line}")
            print()

    print(f" {'=' * 70}")
    print(f" 翻译要求（德语原文 → 中文译文为优先，英文仅供参考）：")
    print(f"  1. 首要目标：德语原文与中文译文之间的准确对应")
    print(f"  2. 宗教术语、圣经引用严格对齐中文和合本 (CUV)")
    print(f"  3. 语义准确优先于文学修饰")
    print(f"  4. 保持诗节/诗行结构")
    print(f"  5. 英文译文仅作语义参考，不体现在最终 docx 文档中")
    print(f"  6. 翻译完成后调用 write_chorale_translations() 回写 Docx")
    print(f" {'=' * 70}")


# ═══════════════════════════════════════════════════════════════
# Write-back: replace 【待翻译】 in docx with Chinese text
# ═══════════════════════════════════════════════════════════════

def write_chorale_translations(chorale_id, translations, verse_translations=None):
    """Write Chinese translations into a chorale .docx file.

    Replaces 【待翻译】 placeholder paragraphs with the provided
    Chinese text, preserving font conventions (宋体, black, 11pt).

    Args:
        chorale_id: str, e.g., "Chorale012"
        translations: list of str — one Chinese text per line/placeholder,
                      in the same order as they appear in the docx.
        verse_translations: optional dict {verse_num (int or str): [chinese_lines]}
            When provided, also persists translations to ChoraleNNN.json
            under the 'chinese_text' field for reuse by the cantata pipeline.

    Returns:
        int: number of placeholders replaced
    """
    docx_path = os.path.join(cfg.DOCX_DIR, f'{chorale_id}_德中对照译文.docx')
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Docx not found: {docx_path}")

    doc = Document(docx_path)
    total_ph = sum(1 for p in doc.paragraphs if '【待翻译】' in p.text)

    # ── Enforce 1:1 alignment ──
    if len(translations) != total_ph:
        print(f"  [WARN] write_chorale_translations: {len(translations)} translations "
              f"but {total_ph} placeholders — padding/truncating to match")
        if len(translations) < total_ph:
            translations = list(translations) + ['[行缺]'] * (total_ph - len(translations))
        else:
            translations = translations[:total_ph]

    trans_idx = 0
    replaced = 0

    for p in doc.paragraphs:
        if '【待翻译】' in p.text and trans_idx < len(translations):
            cn_text = translations[trans_idx]
            # Clear all runs
            for run in p.runs:
                run.text = ''
            # Write first run with Chinese text + correct font
            if p.runs:
                r = p.runs[0]
                r.text = cn_text
                r.font.size = Pt(11)
                r.font.name = 'Times New Roman'
                r.font.color.rgb = RGBColor(0, 0, 0)  # pure black
                r.bold = False
                r.italic = False
                # Set East-Asian font
                rPr = r._r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    from docx.oxml import OxmlElement
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')
            replaced += 1
            trans_idx += 1

    doc.save(docx_path)

    # ── Persist to JSON for cantata pipeline reuse ──
    if verse_translations:
        chorale_json_path = os.path.join(cfg.DATA_DIR, f'{chorale_id}.json')
        if os.path.exists(chorale_json_path):
            with open(chorale_json_path, 'r', encoding='utf-8') as f:
                chorale_data = json.load(f)

            chinese_text = {}
            for vnum, cn_lines in verse_translations.items():
                chinese_text[str(vnum)] = {
                    'lines': cn_lines,
                    'translator': 'AI (manual review)',
                    'translated_at': datetime.now().strftime('%Y-%m-%dT%H:%M:%S'),
                }

            chorale_data['chinese_text'] = chinese_text

            with open(chorale_json_path, 'w', encoding='utf-8') as f:
                json.dump(chorale_data, f, ensure_ascii=False, indent=2)
            print(f"  [JSON] Saved {len(chinese_text)} translations to {chorale_json_path}")

    return replaced


# ═══════════════════════════════════════════════════════════════
# Convenience: full pipeline per BWV
# ═══════════════════════════════════════════════════════════════

def ensure_chorale_resources(bwv_number):
    """Ensure chorale data and docx exist for a BWV, scraping if needed.

    Use this before translate_chorale if you want automatic setup.

    Returns:
        list of chorale IDs with ready docx files
    """
    bwv_str = str(bwv_number)
    entries = idx.lookup_by_bwv(bwv_str)

    ready = []
    for entry in entries:
        cid = entry.get('chorale_id') or entry.get('id')
        if not cid:
            continue

        # Scrape if needed
        chorale_data = scraper.load_chorale_data(cid)
        if not chorale_data:
            chorale_data = scraper.scrape_chorale_detail(cid)
            scraper.save_chorale_data(cid, chorale_data)

        # Generate docx if needed
        docx_path = os.path.join(cfg.DOCX_DIR, f'{cid}_德中对照译文.docx')
        if not os.path.exists(docx_path):
            from . import chorale_translator
            chorale_translator.generate_chorale_docx(chorale_data, chorale_id=cid)

        ready.append(cid)

    return ready


def sync_chorale_glossary(chorale_data, bwv_number):
    """Extract key religious terms from a translated chorale and add to glossary.

    Called after write_chorale_translations() completes, this analyzes the
    German text and Chinese translations for religious terms, then calls
    the pipeline glossary DB to persist them with BWV annotation.

    Args:
        chorale_data: dict with chinese_text, german_text fields
        bwv_number: str — BWV number for annotation
    """
    chinese_text = chorale_data.get('chinese_text', {})
    german_text = chorale_data.get('german_text', {})

    # Common religious terms to detect (German → likely in chorale texts)
    RELIGIOUS_TERMS = {
        'herr': '主', 'gott': '上帝', 'jesus': '耶稣', 'christus': '基督',
        'sünde': '罪', 'gnade': '恩典', 'glaube': '信', 'glauben': '信心',
        'geist': '灵', 'heilig': '圣', 'vater': '父', 'sohn': '子',
        'himmel': '天', 'reich': '国', 'ewig': '永', 'tod': '死',
        'leben': '生命', 'kreuz': '十字架', 'blut': '血', 'lamm': '羔羊',
        'halleluja': '哈利路亚', 'amen': '阿们', 'wort': '道',
        'gerechtigkeit': '公义', 'friede': '和平', 'freude': '喜乐',
        'barmherzigkeit': '怜悯', 'erbarmen': '怜恤', 'liebe': '爱',
        'lob': '赞美', 'dank': '感谢', 'preis': '颂赞',
        'seele': '灵魂', 'herz': '心', 'selig': '福',
        'auferstehung': '复活', 'auferstanden': '复活',
    }

    entries = []
    seen = set()

    for vnum, vdata in german_text.items():
        cn_data = chinese_text.get(str(vnum))
        if not cn_data:
            continue

        cn_lines = cn_data.get('lines', [])
        de_lines = vdata.get('lines', [])

        for di, de_line in enumerate(de_lines):
            de_lower = de_line.lower()
            for term_de, term_cn_hint in RELIGIOUS_TERMS.items():
                if term_de in de_lower and term_de not in seen:
                    seen.add(term_de)
                    # Try to find the actual Chinese translation
                    cn_term = term_cn_hint
                    if di < len(cn_lines):
                        cn_line = cn_lines[di]
                        # Simple heuristic: use the hint or detect from context
                        cn_term = term_cn_hint

                    entries.append({
                        'german': term_de,
                        'chinese_cuv': cn_term,
                        'note': f'[众赞歌] {chorale_data.get("title", "")[:30]} 第{vnum}节',
                    })

    if entries:
        sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))
        try:
            from pipeline import glossary_db
            result = glossary_db.update_from_glossary(str(bwv_number), entries)
            print(f"  [术语库] 已收录 {result.get('new', 0)} 个新术语, "
                  f"{result.get('updated', 0)} 个已更新 (共 {result.get('total', 0)} 条)")
        except Exception as e:
            print(f"  [术语库] 收录失败: {e}")

    return entries
