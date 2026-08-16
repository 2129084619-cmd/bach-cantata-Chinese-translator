# -*- coding: utf-8 -*-
"""Chorale Docx generator — creates German-Chinese parallel translation .docx files.

Font conventions (identical to main cantata docx):
  - Title / metadata labels: Times New Roman bold
  - German text: Times New Roman regular (not bold)
  - Chinese translation: 宋体 regular, black (not gray)
  - Footer / source info: small gray italic
"""

import os
import shutil
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import chorale_config as cfg


def _set_font(run, font_name='Times New Roman', east_asian=None, bold=None, italic=None):
    """Ensure consistent font settings on a run."""
    run.font.name = font_name
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if east_asian:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), east_asian)


def _p(doc, text, bold=False, italic=False, size=None, color=None,
       align=None, sa=Pt(6), left_indent=None):
    """Add a single-run paragraph with consistent font."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    _set_font(r, 'Times New Roman', east_asian='\u5b8b\u4f53')
    if size:
        r.font.size = size
    if color:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    p.paragraph_format.space_after = sa
    return p


def _h(doc, text, level=2):
    """Add a heading with consistent font."""
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        _set_font(r, 'Times New Roman', east_asian='\u5b8b\u4f53')


def _hr(doc):
    """Add a horizontal rule (bottom-border paragraph)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '6',
        qn('w:space'): '1', qn('w:color'): 'AAAAAA'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _gap(doc):
    """Add empty line."""
    doc.add_paragraph()


def _de_line(doc, text):
    """German text line — TNR regular, 11pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.size = Pt(11)
    _set_font(r, 'Times New Roman', bold=False)
    return p


def _cn_line(doc, text=None):
    """Chinese translation line — 宋体 regular, 11pt, black."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    display = text if text else '\u3010\u5f85\u7ffb\u8bd1\u3011'
    r = p.add_run(display)
    r.font.size = Pt(11)
    _set_font(r, 'Times New Roman', east_asian='\u5b8b\u4f53')
    return p


def generate_chorale_docx(chorale_data, chorale_id=None):
    """Generate a chorale .docx file with German-Chinese parallel translation.

    Args:
        chorale_data: dict from scraper (or loaded from JSON)
        chorale_id: optional chorale ID for filename

    Returns:
        str: path to the generated .docx file
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')

    title = chorale_data.get('title', 'Unknown Chorale')
    chorale_id = chorale_id or chorale_data.get('chorale_id', '')

    # ═══════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════
    _p(doc, f'\u201e{title}\u201c',
       bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(2))
    _p(doc, '\u5fb7\u4e2d\u5bf9\u7167\u8bd1\u6587 / German\u2013Chinese Parallel Translation',
       size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    _p(doc, '\u6570\u636e\u6765\u6e90\uff1aBach-Cantatas.com \u4f17\u8d5e\u6b4c\u6587\u672c\u6570\u636e\u5e93',
       italic=True, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
       align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(16))
    _hr(doc)
    _gap(doc)

    # ═══════════════════════════════════════════════════════════════
    # BASIC INFORMATION TABLE
    # ═══════════════════════════════════════════════════════════════
    _h(doc, '\u57fa\u672c\u4fe1\u606f / General Information', level=2)

    info_items = [('\u4f17\u8d5e\u6b4c / Chorale', title)]

    author = chorale_data.get('author', '')
    author_year = chorale_data.get('author_year', '')
    if author:
        label = '\u4f5c\u8005 / Author'
        value = f'{author} ({author_year})' if author_year else author
        info_items.append((label, value))

    melody = chorale_data.get('melody', '')
    if melody:
        info_items.append(('\u65cb\u5f8b / Melody', melody))

    composer = chorale_data.get('composer', '')
    composer_year = chorale_data.get('composer_year', '')
    if composer:
        value = f'{composer} ({composer_year})' if composer_year else composer
        info_items.append(('\u4f5c\u66f2\u5bb6 / Composer', value))

    ekg = chorale_data.get('ekg', '')
    if ekg:
        info_items.append(('EKG \u7f16\u53f7 / EKG', ekg))

    theme = chorale_data.get('theme', '')
    if theme:
        info_items.append(('\u4e3b\u9898 / Theme', theme))

    description = chorale_data.get('description', '')
    if description:
        # Truncate long descriptions to a reasonable length
        desc_short = description[:200] + ('...' if len(description) > 200 else '')
        info_items.append(('\u7b80\u4ecb / Description', desc_short))

    # BWV usage summary
    vocal_works = chorale_data.get('vocal_works', [])
    bach_verses = chorale_data.get('bach_verses', [])
    german_text = chorale_data.get('german_text', {})

    if bach_verses:
        info_items.append((
            '\u5df4\u8d6b\u91c7\u7528\u8bd7\u8282 / Verses set by Bach',
            ', '.join(str(v) for v in bach_verses)
        ))

    info_table = doc.add_table(rows=len(info_items), cols=2, style='Light Grid Accent 1')
    for i, (k, v) in enumerate(info_items):
        info_table.cell(i, 0).text = k
        info_table.cell(i, 1).text = v
        for ci in [0, 1]:
            for para in info_table.cell(i, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    for row in info_table.rows:
        row.cells[0].width = Cm(5.5)

    _gap(doc)

    # ═══════════════════════════════════════════════════════════════
    # VOCAL WORKS TABLE
    # ═══════════════════════════════════════════════════════════════
    if vocal_works:
        _h(doc, '\u5df4\u8d6b\u58f0\u4e50\u4f5c\u54c1\u4f7f\u7528 / Vocal Works by J.S. Bach', level=2)

        # Table: Verse | BWV | Mvt. | Year | BC | Type
        headers = ['Verse', 'BWV', 'Mvt.', 'Year', 'BC', 'Type']
        vw_table = doc.add_table(rows=len(vocal_works) + 1, cols=6, style='Light Grid Accent 1')

        # Header row
        for j, h in enumerate(headers):
            cell = vw_table.cell(0, j)
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for i, vw in enumerate(vocal_works):
            row_data = [
                vw.get('verse', ''),
                vw.get('bwv', ''),
                vw.get('movement', ''),
                vw.get('year', ''),
                vw.get('bc', ''),
                vw.get('type', ''),
            ]
            for j, val in enumerate(row_data):
                vw_table.cell(i + 1, j).text = val
                for para in vw_table.cell(i + 1, j).paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        _gap(doc)

    _hr(doc)
    _gap(doc)

    # ═══════════════════════════════════════════════════════════════
    # GERMAN TEXT — verse by verse, line by line
    # ═══════════════════════════════════════════════════════════════
    _h(doc, '\u5fb7\u8bed\u539f\u6587 / German Text', level=1)

    if bach_verses:
        bach_note = ', '.join(str(v) for v in sorted(bach_verses, key=int))
        _p(doc, f'\u5df4\u8d6b\u91c7\u7528\u8bd7\u8282\uff1a\u7b2c {bach_note} \u8282 (\u52a0\u7c97\u6807\u6ce8)',
           italic=True, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99), sa=Pt(8))

    for verse_num in sorted(german_text.keys(), key=int):
        verse_data = german_text[verse_num]
        lines = verse_data.get('lines', [])
        is_bach = verse_data.get('bold', False)

        # Verse number header
        if is_bach:
            _p(doc, f'{verse_num}.', bold=True, size=Pt(10), sa=Pt(2), left_indent=Cm(0.3))
        else:
            _p(doc, f'{verse_num}.', size=Pt(10), sa=Pt(2), left_indent=Cm(0.3),
               color=RGBColor(0x88, 0x88, 0x88))

        for line in lines:
            _de_line(doc, line)
            # One Chinese placeholder per German line (not per verse)
            _cn_line(doc)

        _gap(doc)

    _hr(doc)

    # ═══════════════════════════════════════════════════════════════
    # FOOTER
    # ═══════════════════════════════════════════════════════════════
    _gap(doc)
    _p(doc, f'\u672c\u6587\u6863\u7531 Chorale Translation Pipeline \u4e8e '
       f'{datetime.now().strftime("%Y-%m-%d")} \u81ea\u52a8\u751f\u6210\uff0c'
       f'\u4e2d\u6587\u7ffb\u8bd1\u5f85\u5b8c\u6210\u3002',
       italic=True, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
       align=WD_ALIGN_PARAGRAPH.CENTER)

    # Save — archive any existing docx first so a re-translation never
    # destroys the prior completed version (mirrors the main cantata pipeline).
    output_path = os.path.join(
        cfg.DOCX_DIR,
        f'{chorale_id}_\u5fb7\u4e2d\u5bf9\u7167\u8bd1\u6587.docx'
    )
    _archive_existing_docx(chorale_id, output_path)
    doc.save(output_path)
    print(f"  [DOCX] Saved: {output_path}")
    return output_path


def _archive_existing_docx(chorale_id, output_path):
    """Archive a previously completed chorale docx before overwriting it.

    Moves the prior version into `translation archive/<ChoraleNNN>/` with a
    timestamp suffix, preserving the full history across re-translations.
    """
    if not os.path.exists(output_path):
        return
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    archive_dir = os.path.join(cfg.ARCHIVE_DIR, chorale_id)
    os.makedirs(archive_dir, exist_ok=True)
    base, ext = os.path.splitext(os.path.basename(output_path))
    backup_path = os.path.join(archive_dir, f'{base}_{ts}{ext}')
    shutil.copy2(output_path, backup_path)
    print(f"  [DOCX] Archived prior version → {backup_path}")
